from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Golden_Memories_Implementation_Plan.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table, widths_in):
    total_dxa = sum(int(round(w * 1440)) for w in widths_in)
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_dxa))
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    cols = grid.findall(qn("w:gridCol"))
    for col, width_in in zip(cols, widths_in):
        col.set(qn("w:w"), str(int(round(width_in * 1440))))
    for row in table.rows:
        for cell, width_in in zip(row.cells, widths_in):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(round(width_in * 1440))))
            set_cell_margins(cell)


def style_run(run, font_name="Calibri", size=11, bold=False, color="1A1A1A"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=6, line=1.25, space_rule="auto"):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        style_run(run)


def add_numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Golden Memories Implementation Plan")
style_run(run, size=26, bold=False, color="000000")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_before = Pt(0)
subtitle.paragraph_format.space_after = Pt(8)
subtitle_run = subtitle.add_run("Technical roadmap for the Spring Boot + Thymeleaf MVC product")
style_run(subtitle_run, size=11, bold=False, color="555555")

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.style = "Table Grid"
meta_rows = [
    ("Project", "Ký Ức Vàng"),
    ("Primary stack", "Java Spring Boot, Thymeleaf, HTML/CSS/JavaScript"),
    ("Architecture", "MVC with modular screens and service layers"),
    ("Scope", "Technical product delivery only"),
]
for row, (label, value) in zip(meta.rows, meta_rows):
    row.cells[0].text = label
    row.cells[1].text = value
    set_cell_shading(row.cells[0], "E8EEF5")
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            for r in p.runs:
                style_run(r)
        set_cell_margins(cell)
set_fixed_table_geometry(meta, [1.5, 5.0])
doc.add_paragraph()

p = doc.add_paragraph()
style_paragraph(p, before=0, after=6, line=1.25)
run = p.add_run("Objective")
run.bold = True
style_run(run, size=12, bold=True, color="1F4D78")
run2 = p.add_run(": build the core family-memory platform first, then layer authentication, persistence, editor tools, and publishing flows on top.")
style_run(run2)

p = doc.add_paragraph()
style_paragraph(p, before=0, after=6, line=1.25)
run = p.add_run("Out of scope for this plan")
run.bold = True
style_run(run, size=12, bold=True, color="1F4D78")
run2 = p.add_run(": pricing strategy, marketing campaigns, distribution strategy, and business storytelling content.")
style_run(run2)

doc.add_heading("1. Build Order", level=1)
add_numbered(doc, [
    "Stabilize the authentication foundation with session handling, OAuth2 sign-in, and OTP email verification.",
    "Persist the core product entities: users, parent profiles, projects, stories, photos, draft states, and approvals.",
    "Turn the dashboard into a protected workspace with project-level data loaded from the database.",
    "Add the editor workflow: receive stories, review transcripts, edit memoir content, and track draft revisions.",
    "Add publishing and archive stubs: print handoff, cloud archive, and QR-protected content.",
    "Implement advanced services only after the MVP works end-to-end: image restoration, notifications, and pricing automation.",
])

doc.add_heading("2. Phase Plan", level=1)
phase = doc.add_table(rows=1, cols=4)
phase.style = "Table Grid"
phase.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = phase.rows[0].cells
for idx, text in enumerate(["Phase", "Goal", "Build items", "Exit criteria"]):
    hdr[idx].text = text
    set_cell_shading(hdr[idx], "E8EEF5")
    hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in hdr[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            style_run(r, bold=True)
rows = [
    ("Phase 0", "Foundation", "Project structure, configuration, layout shell, navigation, and base security.", "App compiles and the main pages load consistently."),
    ("Phase 1", "Identity", "Register, login, Facebook OAuth2, OTP email verification, session protection.", "A user can create an account and enter the app securely."),
    ("Phase 2", "Core data", "Project, parent profile, story, photo, and approval entities with CRUD screens.", "The dashboard reflects real persisted data."),
    ("Phase 3", "Workflow", "Editor queue, transcript review, memo editing, revision tracking, and approval states.", "A project can move from stories to approved draft."),
    ("Phase 4", "Publishing", "Print/export handoff, cloud archive, QR content placeholders, completion status.", "Approved projects can be delivered and archived."),
    ("Phase 5", "Enhancements", "AI image restoration integration, reminders, pricing automation, analytics.", "Non-MVP features are layered without breaking the core flow."),
]
for values in rows:
    cells = phase.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].text = value
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[idx].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            for r in p.runs:
                style_run(r)
        if idx == 0:
            set_cell_shading(cells[idx], "F2F4F7")
            for p in cells[idx].paragraphs:
                p.runs[0].bold = True
        set_cell_margins(cells[idx])
set_fixed_table_geometry(phase, [0.75, 1.05, 3.55, 1.15])

doc.add_heading("3. Core Workstreams", level=1)
doc.add_heading("3.1 Authentication", level=2)
add_bullets(doc, [
    "Keep Facebook OAuth2 active and add Google when the client credentials are ready.",
    "Use OTP email as a fallback verification path for users who do not sign in via social login.",
    "Store secrets in environment variables or a secret manager, not in source-controlled properties.",
    "Protect the dashboard and project routes once user sessions are in place.",
])

doc.add_heading("3.2 Product Data Model", level=2)
add_bullets(doc, [
    "User: account identity, login method, role, and contact details.",
    "Parent profile: parent name, relation, Zalo contact, and connection status.",
    "Project: package, timeline, current phase, and approval state.",
    "Story entry: life-stage category, question prompt, voice transcript, and editor notes.",
    "Photo asset: upload metadata, restoration status, and chapter mapping.",
    "Approval record: draft version, reviewer, comments, and final decision.",
])

doc.add_heading("3.3 Screens", level=2)
add_bullets(doc, [
    "Landing page and product overview.",
    "Registration, login, and OTP verification screens.",
    "Parent connection and profile screens.",
    "Project dashboard and progress tracking.",
    "Media vault, story timeline, transcript review, and security panel.",
    "Draft approval and publishing confirmation.",
])

doc.add_heading("3.4 Services", level=2)
add_bullets(doc, [
    "Auth service for login, OTP generation, and social sign-in hooks.",
    "Project service for workflow state and dashboard aggregation.",
    "Story service for prompts, transcripts, and memoir editing flow.",
    "Media service for photos and future image restoration integration.",
    "Publishing service for print/export and archive handoff.",
])

doc.add_heading("4. Risks and Dependencies", level=1)
risk = doc.add_table(rows=1, cols=3)
risk.style = "Table Grid"
risk.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = risk.rows[0].cells
for idx, text in enumerate(["Risk", "Impact", "Mitigation"]):
    hdr[idx].text = text
    set_cell_shading(hdr[idx], "E8EEF5")
    hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in hdr[idx].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            style_run(r, bold=True)
risk_rows = [
    ("Social login keys are delayed", "Login remains partial until Google is ready.", "Keep the local email OTP path functional and isolate provider config."),
    ("Maven plugin drift in offline cache", "Build steps can fail even if code compiles.", "Pin plugin versions and keep the working local settings mirror."),
    ("Scope creep into business features", "The implementation stalls before the product is usable.", "Freeze the MVP around the family workflow and defer pricing/marketing."),
    ("Workflow ambiguity", "Screens may not map cleanly to backend state.", "Use explicit project states and define each transition in the model."),
]
for values in risk_rows:
    cells = risk.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].text = value
        cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[idx].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            for r in p.runs:
                style_run(r)
        set_cell_margins(cells[idx])
set_fixed_table_geometry(risk, [1.5, 2.0, 3.0])

doc.add_heading("5. Recommended Sprint Goal", level=1)
final_p = doc.add_paragraph()
final_p.paragraph_format.space_after = Pt(0)
final_p.paragraph_format.line_spacing = 1.25
run = final_p.add_run("Deliver a protected onboarding-to-dashboard loop.")
style_run(run, bold=True, color="1F3A5F")
tail = final_p.add_run(" That means one user can register, verify by OTP or Facebook, create a parent profile, open the dashboard, and see real persisted project state.")
style_run(tail)

doc.save(OUTPUT)
print(OUTPUT)
